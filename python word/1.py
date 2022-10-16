from pydoc import doc
from docx import Document 
from docx.shared import Inches

document = Document()
document.add_heading("Это будет главный заголовок", 0)

p = document.add_paragraph("Здесь начинается параграф, ")
p.add_run('Тепер текст будет жирным шрифтом, ').bold = True
p.add_run('А эта часть написана курсивом').italic = True

table = document.add_table(rows=1, cols=3)
c = table.rows[0].cells

c[0].text = 'a'
c[1].text = 'b'
c[2].text = 'c'

document.save('test.docx')